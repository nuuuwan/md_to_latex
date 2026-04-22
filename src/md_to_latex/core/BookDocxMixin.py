import re
import subprocess
import sys

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from rich.console import Console

console = Console()

_FONT = "Garamond"
_MAROON = RGBColor(0x80, 0x00, 0x00)
_DARK_GREY = RGBColor(0x50, 0x50, 0x50)


class BookDocxMixin:
    """Mixin for generating DOCX output matching the PDF appearance."""

    # ── Page layout ──────────────────────────────────────────────────────────

    def _docx_setup_page_layout(self, doc):
        """A4 page, 1-inch margins on all sides."""
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # ── Styles ──────────────────────────────────────────────────────────────

    @staticmethod
    def _docx_force_font(style, name, size_pt=None, bold=False, color=None):
        """Set font on a style, clearing theme-font overrides in the XML."""
        if size_pt is not None:
            style.font.size = Pt(size_pt)
        style.font.bold = bold
        if color is not None:
            style.font.color.rgb = color
        # Directly patch w:rFonts to override any asciiTheme attribute
        style_elem = style._element
        rPr = style_elem.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            style_elem.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in (
            "w:asciiTheme",
            "w:hAnsiTheme",
            "w:cstheme",
            "w:eastAsiaTheme",
        ):
            rFonts.attrib.pop(qn(attr), None)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:cs"), name)

    def _docx_setup_styles(self, doc):
        """Apply Garamond font and double spacing to built-in styles."""
        normal = doc.styles["Normal"]
        self._docx_force_font(normal, _FONT, size_pt=12)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        heading_specs = [
            ("Heading 1", 20),
            ("Heading 2", 16),
            ("Heading 3", 14),
        ]
        for style_name, size in heading_specs:
            self._docx_force_font(
                doc.styles[style_name],
                _FONT,
                size_pt=size,
                bold=False,
                color=RGBColor(0, 0, 0),
            )

        for toc_style in ("TOC 1", "TOC 2", "TOC 3"):
            try:
                self._docx_force_font(
                    doc.styles[toc_style], _FONT, size_pt=12
                )
            except KeyError:
                pass

    # ── Running header ──────────────────────────────────────────────────────

    def _docx_setup_header(self, doc):
        """Centred italic book-title header in dark grey (9 pt)."""
        section = doc.sections[0]
        header = section.header
        p = header.paragraphs[0]
        p.clear()
        run = p.add_run(self.title)
        run.font.name = _FONT
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = _DARK_GREY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Front matter ─────────────────────────────────────────────────────────

    def _docx_add_title_page(self, doc):
        """Title, optional subtitle, author, date, word count — all centred."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.title)
        run.font.name = _FONT
        run.font.size = Pt(24)

        if self.subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(self.subtitle)
            run.font.name = _FONT
            run.font.size = Pt(14)

        doc.add_paragraph()  # vertical spacing

        if self.author:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"By {self.author}")
            run.font.name = _FONT
            run.font.size = Pt(18)

        doc.add_paragraph()  # vertical spacing

        if self.year:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(self.year))
            run.font.name = _FONT
            run.font.size = Pt(9)

        if self.word_count:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{self.word_count:,} words")
            run.font.name = _FONT
            run.font.size = Pt(12)

        doc.add_page_break()

    def _docx_add_copyright_page(self, doc):
        """Edition, publisher, and copyright line — all centred."""
        doc.add_paragraph()
        doc.add_paragraph()

        if self.edition:
            p = doc.add_paragraph(self.edition)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.publisher:
            p = doc.add_paragraph(f"Published by {self.publisher}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.year:
            line = f"Copyright \u00a9 {self.year}"
            if self.author:
                line += f" by {self.author}"
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def _docx_add_toc(self, doc):
        """Insert Word's built-in TOC field with a title heading."""
        title_p = doc.add_paragraph()
        r = title_p.add_run("Table of Contents")
        r.font.name = _FONT
        r.font.size = Pt(20)
        r.font.bold = False
        title_p.paragraph_format.space_after = Pt(12)

        # Word TOC field — \o levels, \z hides page numbers in web view,
        # \u uses applied paragraph outline levels. No \h to avoid the
        # "fields that may refer to other files" security prompt.
        p = doc.add_paragraph()
        run = p.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-3" \\z \\u '

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)

        doc.add_page_break()

    # ── Footnotes ───────────────────────────────────────────────────────────

    def _docx_flush_notes(self, doc):
        """Emit collected footnote texts as end-notes, then reset the list."""
        if not self._fn_notes:
            return
        sep = doc.add_paragraph()
        sep.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        sep.paragraph_format.space_before = Pt(12)
        sep_r = sep.add_run("─" * 30)
        sep_r.font.name = _FONT
        sep_r.font.size = Pt(8)
        sep_r.font.color.rgb = _DARK_GREY
        for num, text in self._fn_notes:
            note_p = doc.add_paragraph()
            note_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            ref_r = note_p.add_run(str(num))
            ref_r.font.name = _FONT
            ref_r.font.size = Pt(8)
            ref_r.font.superscript = True
            text_r = note_p.add_run(f" {text}")
            text_r.font.name = _FONT
            text_r.font.size = Pt(9)
        self._fn_notes = []

    # ── Inline markdown parser ──────────────────────────────────────────────

    def _docx_parse_inline(self, doc, paragraph, text):
        """
        Parse inline markdown and append formatted runs to *paragraph*.

        Handles: ***bold-italic***, **bold**, *italic*, "quoted" (maroon),
        ^[footnote] (superscript number, text collected for end-notes), plain text.
        """
        pattern = re.compile(
            r"(\*\*\*(.+?)\*\*\*"
            r"|\*\*(.+?)\*\*"
            r"|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)"
            r'|"([^"]+)"'
            r"|\^\[([^\]]+)\]"
            r"|([^*\"^]+|\^(?!\[)))",
            re.DOTALL,
        )
        for m in pattern.finditer(text):
            if m.group(2):  # ***bold-italic***
                run = paragraph.add_run(m.group(2))
                run.bold = True
                run.italic = True
            elif m.group(3):  # **bold**
                run = paragraph.add_run(m.group(3))
                run.bold = True
            elif m.group(4):  # *italic*
                run = paragraph.add_run(m.group(4))
                run.italic = True
            elif m.group(5):  # "quoted" → maroon curly quotes
                run = paragraph.add_run(f"\u201c{m.group(5)}\u201d")
                run.font.color.rgb = _MAROON
            elif m.group(6):  # ^[footnote] → superscript number + collect
                self._fn_counter += 1
                self._fn_notes.append((self._fn_counter, m.group(6)))
                ref_r = paragraph.add_run(str(self._fn_counter))
                ref_r.font.name = _FONT
                ref_r.font.size = Pt(8)
                ref_r.font.superscript = True
            elif m.group(7):  # plain text
                paragraph.add_run(m.group(7))

    # ── Block-level content ─────────────────────────────────────────────────

    def _docx_add_markdown_content(self, doc, content):
        """
        Convert markdown body text to docx paragraphs.

        Handles headings, scene breaks (--- / ...), and inline formatting.
        Double-newlines delimit block boundaries.
        """
        blocks = re.split(r"\n\n+", content.strip())

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            # Scene-break block (whole block is --- or ...)
            if re.fullmatch(r"(---|\.\.\.)", block):
                p = doc.add_paragraph("\u2026")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # Heading block
            m = re.match(r"^(#{1,4})\s+(.+)", block)
            if m:
                level = min(len(m.group(1)), 3)
                doc.add_heading(m.group(2).strip(), level=level)
                continue

            # Paragraph (may contain soft-wrapped lines or embedded scene
            # breaks)
            lines = block.split("\n")
            current_para = None

            for line in lines:
                if re.fullmatch(r"\s*(---|\.\.\.)\s*", line):
                    current_para = None
                    p = doc.add_paragraph("\u2026")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    if current_para is None:
                        current_para = doc.add_paragraph()
                    else:
                        current_para.add_run(" ")
                    self._docx_parse_inline(doc, current_para, line.strip())

    def _docx_add_about_section(self, doc, title, content):
        """Add an about-the-book / about-the-author section."""
        doc.add_page_break()
        doc.add_heading(title, level=1)
        self._docx_add_markdown_content(doc, content)
        self._docx_flush_notes(doc)

    # ── Main entry point ────────────────────────────────────────────────────

    def _generate_docx(self, output_path):
        """Build and save the DOCX file to *output_path*.docx."""
        # Reset footnote state
        self._fn_notes = []
        self._fn_counter = 0

        doc = DocxDocument()

        self._docx_setup_page_layout(doc)
        self._docx_setup_styles(doc)
        self._docx_setup_header(doc)

        # Front matter
        self._docx_add_title_page(doc)
        self._docx_add_copyright_page(doc)
        self._docx_add_toc(doc)

        # About sections (mirrors PDF front matter order)
        if self.about_book:
            self._docx_add_about_section(
                doc,
                self.about_book_title or "About The Book",
                self.about_book,
            )
        if self.about_author:
            self._docx_add_about_section(
                doc,
                self.about_author_title or "About the Author",
                self.about_author,
            )

        # Body
        if self.format == 2:
            for chapter in self.chapters:
                doc.add_page_break()
                doc.add_heading(chapter.title, level=1)
                content = chapter._strip_first_heading(chapter.content)
                self._docx_add_markdown_content(doc, content)
                self._docx_flush_notes(doc)
        else:
            for part in self.parts:
                doc.add_page_break()
                doc.add_heading(part.title, level=1)
                for chapter in part.chapters:
                    doc.add_page_break()
                    doc.add_heading(chapter.title, level=2)
                    content = chapter._strip_first_heading(chapter.content)
                    self._docx_add_markdown_content(doc, content)
                    self._docx_flush_notes(doc)

        docx_path = f"{output_path}.docx"
        doc.save(docx_path)
        console.print(
            f"[green]✓ DOCX generated successfully:[/green] "
            f"[bold]{docx_path}[/bold]"
        )

        if sys.platform == "darwin":
            subprocess.run(["open", docx_path], check=False)

        return docx_path
